"""
AvionPure Website — Word Documentation Builder
Generates docs/AvionPure-Website-Documentation.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ──────────────────────────────────────────────
CYAN   = RGBColor(0x00, 0x95, 0xBB)
PURPLE = RGBColor(0x6D, 0x28, 0xD9)
DARK   = RGBColor(0x11, 0x18, 0x27)
MID    = RGBColor(0x4B, 0x55, 0x63)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF0, 0xF4, 0xF8)
BGRAY  = RGBColor(0xE2, 0xE8, 0xF0)

doc = Document()

# ── Page margins ───────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ─────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Helper: shade table cell ───────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: set cell borders ───────────────────────────────────
def set_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'E2E8F0')
        borders.append(b)
    tblPr.append(borders)

# ── Helper: add a heading ──────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    shade_paragraph(p, '111827')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.1)
    run = p.add_run(text)
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = WHITE
    run.font.name  = 'Calibri'
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = CYAN
    run.font.name  = 'Calibri'
    # Bottom border via paragraph border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '0095BB')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = DARK
    run.font.name  = 'Calibri'
    return p

def body(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.name  = 'Calibri'
    run.font.color.rgb = color if color else MID
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = MID
    return p

def label(text, value, label_color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{text}: ")
    r1.font.bold  = True
    r1.font.size  = Pt(10)
    r1.font.color.rgb = label_color if label_color else DARK
    r1.font.name  = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size  = Pt(10)
    r2.font.color.rgb = MID
    r2.font.name  = 'Calibri'

def code_line(text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'F0F4F8')
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = PURPLE
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, '111827')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run  = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name  = 'Calibri'

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = 'F7F9FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p    = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run  = p.add_run(str(cell_text))
            run.font.size  = Pt(9)
            run.font.color.rgb = MID
            run.font.name  = 'Calibri'

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return table

def callout(text, color='0095BB', bg='EBF7FB'):
    p = doc.add_paragraph()
    shade_paragraph(p, bg)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor.from_string(color)
    return p

# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
shade_paragraph(p, '111827')
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AvionPure')
r.font.size  = Pt(36)
r.font.bold  = True
r.font.color.rgb = WHITE
r.font.name  = 'Calibri'

p2 = doc.add_paragraph()
shade_paragraph(p2, '111827')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(0)
r2 = p2.add_run('Transforming Innovation')
r2.font.size  = Pt(14)
r2.font.color.rgb = CYAN
r2.font.name  = 'Calibri'

p3 = doc.add_paragraph()
shade_paragraph(p3, '111827')
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(16)
p3.paragraph_format.space_after  = Pt(0)
r3 = p3.add_run('Website Documentation')
r3.font.size  = Pt(22)
r3.font.bold  = True
r3.font.color.rgb = RGBColor(0xA0, 0xB4, 0xC8)
r3.font.name  = 'Calibri'

for _ in range(3):
    p = doc.add_paragraph()
    shade_paragraph(p, '111827')
    p.paragraph_format.space_after = Pt(0)

meta_items = [
    ('Version',      'v2.0  —  Light Theme'),
    ('Project',      'prj-05'),
    ('Last Updated', 'March 2026'),
    ('Pages',        'index.html  ·  apply.html'),
    ('Contact',      'ask@avionpure.com'),
    ('Location',     'Omaha, NE. USA'),
]
for k, v in meta_items:
    p = doc.add_paragraph()
    shade_paragraph(p, '111827')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{k}:  ')
    r1.font.bold  = True
    r1.font.size  = Pt(10)
    r1.font.color.rgb = RGBColor(0x8B, 0xA5, 0xC2)
    r1.font.name  = 'Calibri'
    r2 = p.add_run(v)
    r2.font.size  = Pt(10)
    r2.font.color.rgb = WHITE
    r2.font.name  = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════
h1('1.  Project Overview')
body('AvionPure is an IT solutions company specialising in Security & Observability, '
     'Cloud Infrastructure, Software Development, and Technology Consulting. '
     'The website is a corporate marketing site with a linked job application page. '
     'It is built with vanilla HTML, CSS, and JavaScript — no frameworks or build tools required.')

h3('Key Facts')
label('Website Type',   'Corporate single-page site + job application page')
label('Theme',          'Light / white background with cyan (#0095bb) and purple (#6d28d9) brand accents')
label('Responsive',     'Yes — breakpoints at 1024px, 860px, 768px, 600px, 480px')
label('Dependencies',   'Google Fonts only (Inter + Space Grotesk). No npm, no bundler.')
label('Hosting',        'Static HTML — can deploy to Netlify, GitHub Pages, AWS S3, or any web server')

divider()

# ══════════════════════════════════════════════════════════════
#  2. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════
h1('2.  File Structure')
body('All website files live in prj-05/. The docs/ folder contains documentation only.')

h3('Directory Layout')
for line in [
    'prj-05/',
    '  ├── index.html         ← Main website (single-page)',
    '  ├── apply.html         ← Job application page (URL-param driven)',
    '  ├── styles.css         ← All styles for both pages',
    '  ├── script.js          ← All JavaScript behaviours',
    '  └── docs/',
    '        └── documentation.html   ← HTML documentation',
    '        └── AvionPure-Website-Documentation.docx  ← This file',
]:
    code_line(line)

h3('File Reference')
add_table(
    ['File', 'Purpose', 'Size'],
    [
        ['index.html',   'Main corporate website — all sections in one page', '~22 KB'],
        ['apply.html',   'Job application page — reads ?job= URL param, renders job details + form', '~28 KB'],
        ['styles.css',   'Shared stylesheet for both pages — light theme, CSS variables, responsive', '~23 KB'],
        ['script.js',    'Navbar scroll, hamburger, particle canvas, scroll-reveal, counter, contact form', '~6 KB'],
    ],
    col_widths=[1.6, 4.0, 0.9]
)

divider()

# ══════════════════════════════════════════════════════════════
#  3. BRAND & THEME
# ══════════════════════════════════════════════════════════════
h1('3.  Brand & Theme')

h2('3.1  Company Identity')
add_table(
    ['Element', 'Value'],
    [
        ['Company Name', 'AvionPure  —  "Avion" in dark, "Pure" in cyan'],
        ['Tagline',      'Transforming Innovation  (cyan, below logo name)'],
        ['Positioning',  'IT Solutions & Modernization'],
        ['Email',        'ask@avionpure.com'],
        ['Location',     'Omaha, NE. USA'],
    ],
    col_widths=[2.0, 4.5]
)

h2('3.2  Colour Palette')
add_table(
    ['Colour', 'Hex Code', 'Usage'],
    [
        ['Cyan (Primary)',       '#0095bb', 'Logo "Pure", tagline, links, tags, accents, gradient start'],
        ['Purple (Accent)',      '#6d28d9', 'Gradient end, featured badges, alternate accents'],
        ['Dark (Text/Headers)',  '#111827', 'Headings, navbar background, dark panels'],
        ['Mid (Body Text)',      '#4b5563', 'Paragraphs, descriptions, secondary text'],
        ['Light BG',            '#f7f9fc', 'Alternate section backgrounds, card surfaces'],
        ['Border',              '#e2e8f0', 'Card borders, dividers, separators'],
        ['Gradient',            '#0095bb → #6d28d9', 'Buttons, badges, active highlights'],
    ],
    col_widths=[1.8, 1.4, 3.4]
)

h2('3.3  Typography')
add_table(
    ['Font', 'Weights Used', 'Applied To'],
    [
        ['Space Grotesk', '400, 500, 600, 700', 'Page headings (h1–h2), logo name, section titles'],
        ['Inter',         '300, 400, 500, 600, 700, 800, 900', 'Body text, nav links, buttons, labels, paragraphs'],
    ],
    col_widths=[1.8, 1.8, 3.0]
)

divider()

# ══════════════════════════════════════════════════════════════
#  4. PAGES
# ══════════════════════════════════════════════════════════════
h1('4.  Pages')

h2('4.1  index.html  —  Main Website')
body('A single-page website with smooth-scroll navigation. All sections use anchor IDs. '
     'The fixed navbar highlights the active section as the user scrolls.')

h3('Page Flow')
for step in ['#navbar  →  #hero  →  #services  →  #about  →  #careers  →  #contact  →  <footer>']:
    code_line(step)

h2('4.2  apply.html  —  Job Application Page')
body('Opens in a new browser tab when "Apply Now" is clicked on any job card. '
     'Reads the ?job= URL parameter and dynamically renders the job details and application form using JavaScript.')

h3('URL Parameters')
add_table(
    ['URL', 'Job Loaded'],
    [
        ['apply.html?job=sre-engineer',            'SRE Engineer'],
        ['apply.html?job=siem-architect',           'SIEM Architect'],
        ['apply.html?job=observability-engineer',  'Observability Engineer'],
        ['apply.html?job=security-engineer',        'Security Engineer'],
        ['apply.html?job=<anything-else>',          '"Position Not Found" message with link back'],
    ],
    col_widths=[3.2, 3.2]
)

h3('Apply Page Layout — Two Columns')
add_table(
    ['Panel', 'Contents'],
    [
        ['Left — Job Detail (sticky)',
         'Badge, Job Title, Meta tags (type/location/department), Overview, Responsibilities (7 bullets), Requirements (7 bullets), Skill tags'],
        ['Right — Application Form',
         'First Name, Last Name, Email, Phone, LinkedIn, Years Experience (dropdown), Cover Letter (textarea), Resume upload (PDF/DOC/DOCX max 5MB), Submit button → success message'],
    ],
    col_widths=[2.0, 4.5]
)

callout('Form validation: First Name, Email, and Resume are required. '
        'Missing fields trigger a shake animation on the Submit button. '
        'On success the form hides and a confirmation card appears.')

divider()

# ══════════════════════════════════════════════════════════════
#  5. SECTION-BY-SECTION LAYOUT
# ══════════════════════════════════════════════════════════════
h1('5.  Section-by-Section Layout')

# NAV
h2('5.1  Navigation  —  #navbar')
add_table(
    ['Element', 'Detail'],
    [
        ['Position',       'Fixed top, full-width, z-index 1000'],
        ['Default state',  'Transparent — blends with the hero section'],
        ['Scrolled state', 'White background + backdrop blur + border-bottom shadow (triggers at scrollY > 40px)'],
        ['Logo',           '"Avion" (dark) + "Pure" (cyan) + "Transforming Innovation" tagline (cyan, small, below name)'],
        ['Nav Links',      'Services · About · Careers · Contact Us (gradient pill button)'],
        ['Active link',    'Active section link turns cyan via IntersectionObserver'],
        ['Mobile',         'Nav links hidden; hamburger (☰) shown; dropdown panel animates open'],
    ],
    col_widths=[2.0, 4.5]
)

# HERO
h2('5.2  Hero  —  #hero')
add_table(
    ['Element', 'Detail'],
    [
        ['Layout',          'Full-viewport height (100vh). Flex row: content left, orbit visual right.'],
        ['Background',      'White with subtle radial gradients (purple top-centre, cyan right)'],
        ['Particle Canvas', '80 animated particles (cyan + purple, low opacity). Particles within 120px are connected by a line. Resizes on window resize.'],
        ['Main Title',      '"IT Solutions &" (dark text) + "Modernization" (cyan-purple gradient text)'],
        ['Subtitle',        'One-sentence brand description paragraph'],
        ['CTA Buttons',     '"Explore Services" (gradient pill)  ·  "Get in Touch" (ghost outline pill)'],
        ['Orbit Visual',    '3 animated rings + floating centre logo + 3 orbiting dots. Visible on screens ≥ 1024px only.'],
        ['Scroll Indicator','Animated vertical line + "Scroll" label — bottom-centre of section'],
    ],
    col_widths=[2.0, 4.5]
)
callout('Note: The hero stats block (150+ clients, 98% uptime, 12+ years) was removed per feedback. '
        'Counter animation code remains in script.js but is inactive — the observer checks gracefully.')

# SERVICES
h2('5.3  Services  —  #services')
body('4-column horizontal grid on desktop. Reduces to 2-col at ≤1024px and 1-col at ≤600px.')
add_table(
    ['#', 'Title', 'Skill Tags', 'Badge'],
    [
        ['1', 'Security & Observability',  'SIEM, ITSI, Datadog, Cribl',               '—'],
        ['2', 'IT Infra & Cloud Solutions','AWS, Azure, GCP, Kubernetes',               'Popular'],
        ['3', 'Software Development',      'APIs, Microservices, Automation, DevOps',   '—'],
        ['4', 'Technology Consulting',     'Strategy, Architecture, Roadmap, Advisory', '—'],
    ],
    col_widths=[0.3, 2.0, 2.8, 1.0]
)
body('Cards animate up on scroll (IntersectionObserver). Hover lifts 5px with cyan border glow. "Learn more" links removed per feedback.')

# ABOUT
h2('5.4  About  —  #about')
add_table(
    ['Element', 'Detail'],
    [
        ['Layout',     'Single-column, centred, max-width 820px. Tech badges grid removed per feedback.'],
        ['Background', 'Light gray (#f7f9fc)'],
        ['Section Tag','Why AvionPure'],
        ['Title',      '"Built for the Future of IT"'],
        ['Pillars',    '3-column card row: Innovation First  ·  Security by Design  ·  Measurable Outcomes'],
    ],
    col_widths=[2.0, 4.5]
)

# CAREERS
h2('5.5  Careers  —  #careers')
body('4-column horizontal grid on desktop. Reduces to 2-col at ≤1024px and 1-col at ≤600px.')
add_table(
    ['Job Title', 'Type', 'Badge', 'Apply Link'],
    [
        ['SRE Engineer',          'Full-time · Remote', '—',           'apply.html?job=sre-engineer'],
        ['SIEM Architect',        'Full-time · Hybrid', '—',           'apply.html?job=siem-architect'],
        ['Observability Engineer','Full-time · Remote', 'Hiring Now',  'apply.html?job=observability-engineer'],
        ['Security Engineer',     'Full-time · Hybrid', '—',           'apply.html?job=security-engineer'],
    ],
    col_widths=[1.8, 1.4, 1.0, 2.4]
)
body('"Apply Now" opens apply.html in a new tab (target="_blank") with the relevant ?job= parameter.')

# CONTACT
h2('5.6  Contact  —  #contact')
add_table(
    ['Element', 'Detail'],
    [
        ['Layout',         '2-column: info cards left (1fr) · form right (2fr). Stacks at ≤860px.'],
        ['Background',     'Light gray (#f7f9fc)'],
        ['Contact Cards',  'Email Us (ask@avionpure.com)  ·  Location (Omaha, NE. USA). Website card removed per feedback.'],
        ['Form Fields',    'Full Name · Phone · Email · Subject (dropdown) · Message'],
        ['Validation',     'Name, Email, Message required. Shake animation on missing fields.'],
        ['On Submit',      '1.2s simulated delay → button hides → success banner shown'],
    ],
    col_widths=[2.0, 4.5]
)

# FOOTER
h2('5.7  Footer')
add_table(
    ['Element', 'Detail'],
    [
        ['Layout',          '4-column grid: Brand (2fr) · Services (1fr) · Company (1fr) · Contact (1fr)'],
        ['Brand Column',    'Large logo — icon + "AvionPure" + "Transforming Innovation" tagline. Right-aligned toward columns.'],
        ['Column Headers',  'SERVICES · COMPANY · CONTACT — all in cyan (#0095bb)'],
        ['Services Links',  'All scroll to #services'],
        ['Company Links',   'About · Careers · Contact'],
        ['Contact Info',    'ask@avionpure.com  ·  Omaha, NE. USA'],
        ['Footer Bottom',   '© 2026 AvionPure. All rights reserved.  ·  "Transforming Innovation"'],
    ],
    col_widths=[2.0, 4.5]
)

divider()

# ══════════════════════════════════════════════════════════════
#  6. JAVASCRIPT BEHAVIOURS
# ══════════════════════════════════════════════════════════════
h1('6.  JavaScript Behaviours  (script.js)')

behaviours = [
    ('1 · Navbar Scroll',
     'Listens to window scroll. Adds .scrolled class to #navbar when scrollY > 40. '
     'The .scrolled class applies white background, backdrop-filter blur, and box-shadow.'),
    ('2 · Hamburger Menu',
     'Click on #hamburger toggles .open class on #mobileMenu. '
     '.open animates max-height from 0 to 300px via CSS transition. Any link click closes the menu.'),
    ('3 · Particle Canvas',
     'IIFE renders 80 particles on #particleCanvas using requestAnimationFrame. '
     'Particles drift with random velocity. Particles within 120px are connected with a semi-transparent cyan line. '
     'Canvas resizes on window resize event.'),
    ('4 · Scroll Reveal',
     'IntersectionObserver (threshold 0.1, rootMargin -40px bottom) watches: '
     '.service-card, .job-card, .contact-card, .pillar, .section-header, .about-text. '
     'Adds .reveal (opacity 0, translateY 30px) on load; adds .visible when in viewport. '
     'Staggered delays via inline transitionDelay per element index.'),
    ('5 · Active Nav Highlight',
     'Second IntersectionObserver (threshold 0.4) watches all section[id] elements. '
     'When a section enters viewport its nav link turns cyan; all others reset.'),
    ('6 · Counter Animation (inactive)',
     'Animates .stat-num elements from 0 to data-target value. '
     'Hero stats section was removed from HTML — observer checks if(statsEl) before observing, so fails gracefully.'),
    ('7 · Contact Form',
     'Validates Name, Email, Message. Invalid: shake animation on submit button. '
     'Valid: disables button, sets text to "Sending…", waits 1.2s, hides button, shows #formSuccess banner. '
     'Submission is simulated — no backend connected.'),
]

for title, desc in behaviours:
    h3(title)
    body(desc)

divider()

# ══════════════════════════════════════════════════════════════
#  7. RESPONSIVE DESIGN
# ══════════════════════════════════════════════════════════════
h1('7.  Responsive Design  —  Breakpoints')
add_table(
    ['Breakpoint', 'Layout Changes'],
    [
        ['≥ 1024px',  'Hero orbit visual shown. Services and Careers grids display 4 columns.'],
        ['≤ 1024px',  'Services and Careers grids switch to 2 columns.'],
        ['≤ 860px',   'Hero stacks vertically (no orbit visual). Contact grid stacks. Footer: 2 columns. About pillars: 1 column.'],
        ['≤ 768px',   'Desktop nav links hidden. Hamburger icon shown. Mobile dropdown menu active. Section padding reduces.'],
        ['≤ 600px',   'Services and Careers: 1 column. Form rows stack to single column. Apply page columns stack.'],
    ],
    col_widths=[1.4, 5.1]
)

divider()

# ══════════════════════════════════════════════════════════════
#  8. CSS VARIABLES REFERENCE
# ══════════════════════════════════════════════════════════════
h1('8.  CSS Variables Reference  (styles.css)')
for line in [
    ':root {',
    '  /* Backgrounds */',
    '  --bg:       #ffffff;     /* Page background */',
    '  --bg2:      #f7f9fc;     /* Alternate section background */',
    '  --bg3:      #edf0f7;     /* Input / form field background */',
    '  /* Surfaces */',
    '  --surface:  rgba(0,0,0,0.03);',
    '  --surface2: rgba(0,0,0,0.055);',
    '  /* Borders */',
    '  --border:   rgba(0,0,0,0.09);',
    '  --border2:  rgba(0,120,180,0.25);',
    '  /* Text */',
    '  --text:     #111827;     /* Primary dark text */',
    '  --text2:    #4b5563;     /* Secondary / body text */',
    '  --text3:    #9ca3af;     /* Muted / placeholder */',
    '  /* Brand */',
    '  --cyan:     #0095bb;     /* Primary brand — logo "Pure" colour */',
    '  --purple:   #6d28d9;     /* Accent brand colour */',
    '  --grad:     linear-gradient(135deg, #0095bb, #6d28d9);',
    '  --grad2:    linear-gradient(135deg, #6d28d9, #0095bb);',
    '  /* Geometry */',
    '  --radius:    16px;',
    '  --radius-sm: 10px;',
    '  --transition: 0.3s cubic-bezier(0.4,0,0.2,1);',
    '}',
]:
    code_line(line)

divider()

# ══════════════════════════════════════════════════════════════
#  9. FORMS SUMMARY
# ══════════════════════════════════════════════════════════════
h1('9.  Forms Summary')
add_table(
    ['Form', 'Location', 'Required Fields', 'On Success'],
    [
        ['Contact Form',     'index.html  #contact',  'Full Name, Email, Message',      'Success banner below submit button'],
        ['Application Form', 'apply.html',             'First Name, Email, Resume file', 'Entire form replaced with confirmation card'],
    ],
    col_widths=[1.5, 1.6, 2.0, 1.5]
)
callout('Backend Note: Both forms simulate submission with setTimeout. '
        'To make them functional, connect to Formspree, EmailJS, Netlify Forms, or a custom API endpoint.',
        color='92400E', bg='FEF3C7')

divider()

# ══════════════════════════════════════════════════════════════
#  10. CHANGE LOG
# ══════════════════════════════════════════════════════════════
h1('10.  Change Log')
add_table(
    ['Date', 'Change'],
    [
        ['Mar 2026', 'Migrated all website files from prj-01 to prj-05'],
        ['Mar 2026', 'Renamed styles_v2.css → styles.css; updated references in both HTML files'],
        ['Mar 2026', 'Created apply.html — 4 job descriptions + application form with file upload'],
        ['Mar 2026', 'Removed "Learn more" links from all service cards'],
        ['Mar 2026', 'Updated location to "Omaha, NE. USA" (was "Global — Remote-First")'],
        ['Mar 2026', 'Removed "Website" contact card from contact section'],
        ['Mar 2026', 'Footer column headers (SERVICES, COMPANY, CONTACT) set to cyan'],
        ['Mar 2026', 'Footer brand: logo enlarged, tagline sized up, right-aligned toward columns'],
        ['Mar 2026', 'Hero title changed to "IT Solutions & Modernization"'],
        ['Mar 2026', 'Logo tagline set to lowercase and cyan colour'],
        ['Mar 2026', 'Removed hero stats section (150+ clients, 98% uptime, 12+ years)'],
        ['Mar 2026', 'Removed tech badges grid from About section'],
        ['Mar 2026', 'Services and Careers forced to 4-column horizontal grid'],
        ['Mar 2026', 'Built light theme (styles.css) replacing original dark theme'],
        ['Mar 2026', 'Added logo + tagline to navbar and footer logo'],
        ['Mar 2026', 'Archived original dark-theme index.html as index_1.html in prj-01'],
    ],
    col_widths=[1.2, 5.3]
)

# ══════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════
OUT = '/Users/atulsingh/Documents/VERSION-control-MAC/MY-projects-MAC/AS-Avion/prj-05/docs/AvionPure-Website-Documentation.docx'
doc.save(OUT)
print(f'✓  Saved  →  {OUT}')
